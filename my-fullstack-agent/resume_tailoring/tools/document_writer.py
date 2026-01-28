"""Document writing tool for generating .docx resumes with color-coded content."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from google.adk.tools import FunctionTool
from google.adk.tools.tool_context import ToolContext
from typing import Optional
import json


def _generate_docx(
    name: str,
    email: str,
    phone: str,
    summary: str,
    experiences_json: str,
    skills_json: str,
    linkedin: Optional[str] = None,
    output_path: Optional[str] = None,
    format_json: Optional[str] = None,
    tool_context: Optional[ToolContext] = None,
) -> dict:
    """Generates a professional resume in .docx format with color-coded content.

    Creates a formatted resume document where fabricated/new content is marked
    in red to distinguish it from original content. Optionally applies detected
    formatting from the original resume.

    Args:
        name: Candidate's full name.
        email: Candidate's email address.
        phone: Candidate's phone number.
        summary: Professional summary statement.
        experiences_json: JSON string of experience list, each containing:
            - title: Job title
            - company: Company name
            - date_range: Employment period
            - bullet_points: List of achievement bullet points
            - is_fabricated: Boolean indicating if content is fabricated
        skills_json: JSON string of skills list.
        linkedin: Candidate's LinkedIn URL (optional).
        output_path: Full path where the .docx file should be saved. If not provided,
                     reads from session state.
        format_json: Optional JSON string of format settings containing:
            - body_font_name, body_font_size
            - header_font_name, header_font_size, header_bold
            - section_font_name, section_font_size, section_bold
            - subheader_font_name, subheader_font_size, subheader_bold
            - margin_top, margin_bottom, margin_left, margin_right
        tool_context: The tool context providing access to session state.

    Returns:
        dict: Contains 'status', 'output_path', and optional 'error_message'.
    """
    try:
        # If output_path not provided, read from session state
        if not output_path and tool_context:
            output_path = tool_context.state.get("output_path", "./tailored_resume.docx")
        elif not output_path:
            output_path = "./tailored_resume.docx"
        # Parse JSON inputs
        experiences = json.loads(experiences_json)
        skills = json.loads(skills_json)

        # Parse format settings with defaults
        fmt = {}
        if format_json:
            try:
                fmt = json.loads(format_json)
            except json.JSONDecodeError:
                pass  # Use defaults if format JSON is invalid

        # Default format values
        body_font = fmt.get("body_font_name", "Calibri")
        body_size = fmt.get("body_font_size", 11.0)
        header_font = fmt.get("header_font_name", "Calibri")
        header_size = fmt.get("header_font_size", 24.0)
        header_bold = fmt.get("header_bold", True)
        section_font = fmt.get("section_font_name", "Calibri")
        section_size = fmt.get("section_font_size", 14.0)
        section_bold = fmt.get("section_bold", True)
        subheader_font = fmt.get("subheader_font_name", "Calibri")
        subheader_size = fmt.get("subheader_font_size", 11.0)
        subheader_bold = fmt.get("subheader_bold", True)
        margin_top = fmt.get("margin_top", 1.0)
        margin_bottom = fmt.get("margin_bottom", 1.0)
        margin_left = fmt.get("margin_left", 1.0)
        margin_right = fmt.get("margin_right", 1.0)

        doc = Document()

        # Apply margins
        section = doc.sections[0]
        section.top_margin = Inches(margin_top)
        section.bottom_margin = Inches(margin_bottom)
        section.left_margin = Inches(margin_left)
        section.right_margin = Inches(margin_right)

        # Set up default font style
        style = doc.styles['Normal']
        style.font.name = body_font
        style.font.size = Pt(body_size)

        # Header - Name (custom formatting instead of built-in heading)
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.add_run(name)
        header_run.font.name = header_font
        header_run.font.size = Pt(header_size)
        header_run.font.bold = header_bold

        # Contact Info
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_text = []
        if email:
            contact_text.append(email)
        if phone:
            contact_text.append(phone)
        if linkedin:
            contact_text.append(linkedin)
        contact_run = contact_para.add_run(' | '.join(contact_text))
        contact_run.font.name = body_font
        contact_run.font.size = Pt(body_size)

        # Professional Summary
        summary_header = doc.add_paragraph()
        summary_header_run = summary_header.add_run('Professional Summary')
        summary_header_run.font.name = section_font
        summary_header_run.font.size = Pt(section_size)
        summary_header_run.font.bold = section_bold

        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(summary)
        summary_run.font.name = body_font
        summary_run.font.size = Pt(body_size)

        # Work Experience
        exp_header = doc.add_paragraph()
        exp_header_run = exp_header.add_run('Professional Experience')
        exp_header_run.font.name = section_font
        exp_header_run.font.size = Pt(section_size)
        exp_header_run.font.bold = section_bold

        for exp in experiences:
            is_fabricated = exp.get('is_fabricated', False)

            # Title and Company
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(
                f"{exp['title']} at {exp['company']}"
            )
            title_run.font.name = subheader_font
            title_run.font.size = Pt(subheader_size)
            title_run.font.bold = subheader_bold
            if is_fabricated:
                title_run.font.color.rgb = RGBColor(255, 0, 0)

            # Date range
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(exp.get('date_range', ''))
            date_run.font.name = body_font
            date_run.font.size = Pt(body_size)
            date_run.italic = True
            if is_fabricated:
                date_run.font.color.rgb = RGBColor(255, 0, 0)

            # Bullet points
            for bullet in exp.get('bullet_points', []):
                bullet_para = doc.add_paragraph(style='List Bullet')
                bullet_run = bullet_para.add_run(bullet)
                bullet_run.font.name = body_font
                bullet_run.font.size = Pt(body_size)
                if is_fabricated:
                    bullet_run.font.color.rgb = RGBColor(255, 0, 0)

        # Skills Section
        skills_header = doc.add_paragraph()
        skills_header_run = skills_header.add_run('Skills')
        skills_header_run.font.name = section_font
        skills_header_run.font.size = Pt(section_size)
        skills_header_run.font.bold = section_bold

        skills_para = doc.add_paragraph()
        skills_run = skills_para.add_run(', '.join(skills))
        skills_run.font.name = body_font
        skills_run.font.size = Pt(body_size)

        # Save document
        doc.save(output_path)

        return {
            "status": "success",
            "output_path": output_path,
            "message": f"Resume generated successfully at {output_path}",
            "format_applied": bool(format_json),
        }

    except json.JSONDecodeError as e:
        return {
            "status": "error",
            "output_path": output_path,
            "error_message": f"Invalid JSON input: {str(e)}"
        }
    except Exception as e:
        return {
            "status": "error",
            "output_path": output_path,
            "error_message": str(e)
        }


# Create FunctionTool instance
generate_docx_tool = FunctionTool(func=_generate_docx)


def _tailor_docx_in_place(
    base_resume_path: str,
    edits_json: str,
    output_path: Optional[str] = None,
    tool_context: Optional[ToolContext] = None,
) -> dict:
    """Tailor an existing resume .docx by replacing bullet paragraphs in-place.

    This is the preferred path for strict formatting preservation:
    - no new sections
    - no summary line
    - only edits existing bullets that match exactly

    Args:
      base_resume_path: Path to the existing resume (.docx).
      edits_json: JSON list of BulletEdit-like objects with fields: old_bullet, new_bullet.
      output_path: Output .docx path; defaults to session state's output_path.
      tool_context: ADK tool context.

    Returns:
      status + output_path + counts.
    """
    try:
        if not output_path and tool_context:
            output_path = tool_context.state.get("output_path", "./tailored_resume.docx")
        elif not output_path:
            output_path = "./tailored_resume.docx"

        edits = json.loads(edits_json)
        doc = Document(base_resume_path)

        # Build quick lookup map (normalized text -> list of paragraph indices)
        def norm(s: str) -> str:
            return " ".join((s or "").split()).strip()

        idx_map: dict[str, list[int]] = {}
        for i, p in enumerate(doc.paragraphs):
            t = norm(p.text)
            if not t:
                continue
            idx_map.setdefault(t, []).append(i)

        replaced = 0
        missing = 0
        for e in edits:
            old = norm(e.get("old_bullet", ""))
            new = e.get("new_bullet", "")
            if not old:
                continue
            hits = idx_map.get(old, [])
            if not hits:
                missing += 1
                continue
            # Replace the first occurrence; duplicates are rare but possible.
            p = doc.paragraphs[hits[0]]
            p.text = new
            replaced += 1

        doc.save(output_path)
        return {
            "status": "success",
            "output_path": output_path,
            "replaced": replaced,
            "missing": missing,
        }

    except Exception as e:
        return {
            "status": "error",
            "output_path": output_path,
            "error_message": str(e),
        }


# Tool for in-place tailoring
tailor_docx_in_place_tool = FunctionTool(func=_tailor_docx_in_place)
