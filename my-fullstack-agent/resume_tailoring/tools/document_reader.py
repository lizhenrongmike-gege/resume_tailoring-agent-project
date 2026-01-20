"""Document reading tools for extracting text from .docx and .pdf files."""

from docx import Document
import pdfplumber
from google.adk.tools import FunctionTool


def _read_docx(file_path: str) -> dict:
    """Reads and extracts text content from a .docx file.

    Args:
        file_path: The full path to the .docx file to read.

    Returns:
        dict: Contains 'status', 'file_path', 'content', and optional 'metadata'.
    """
    try:
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        content = "\n".join(paragraphs)
        return {
            "status": "success",
            "file_path": file_path,
            "content": content,
            "metadata": {"paragraph_count": len(doc.paragraphs)}
        }
    except FileNotFoundError:
        return {
            "status": "error",
            "file_path": file_path,
            "error_message": f"File not found: {file_path}"
        }
    except Exception as e:
        return {
            "status": "error",
            "file_path": file_path,
            "error_message": str(e)
        }


def _read_pdf(file_path: str) -> dict:
    """Reads and extracts text content from a .pdf file.

    Args:
        file_path: The full path to the .pdf file to read.

    Returns:
        dict: Contains 'status', 'file_path', 'content', and optional 'metadata'.
    """
    try:
        content_parts = []
        page_count = 0
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    content_parts.append(text)
        content = "\n".join(content_parts)
        return {
            "status": "success",
            "file_path": file_path,
            "content": content,
            "metadata": {"page_count": page_count}
        }
    except FileNotFoundError:
        return {
            "status": "error",
            "file_path": file_path,
            "error_message": f"File not found: {file_path}"
        }
    except Exception as e:
        return {
            "status": "error",
            "file_path": file_path,
            "error_message": str(e)
        }


# Create FunctionTool instances
read_docx_tool = FunctionTool(func=_read_docx)
read_pdf_tool = FunctionTool(func=_read_pdf)
